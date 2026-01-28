import docx
import json
import os

questionnaires_dir = r"C:\Users\rinai\stat-q\src\Questionnaires"
output_dir = r"C:\Users\rinai\stat-q\extracted_questionnaires"

os.makedirs(output_dir, exist_ok=True)

# Get all docx files
docx_files = [f for f in os.listdir(questionnaires_dir) if f.endswith('.docx')]

for docx_file in docx_files:
    file_path = os.path.join(questionnaires_dir, docx_file)

    try:
        doc = docx.Document(file_path)

        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Also extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                tables_text.append(" | ".join(row_text))

        # Save as text file
        output_filename = docx_file.replace('.docx', '.txt')
        output_path = os.path.join(output_dir, output_filename)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("=== PARAGRAPHS ===\n")
            f.write("\n".join(full_text))
            f.write("\n\n=== TABLES ===\n")
            f.write("\n".join(tables_text))

        print(f"[OK] Extracted: {docx_file}")

    except Exception as e:
        print(f"[ERROR] Error extracting {docx_file}: {e}")

print(f"\nAll extracted files saved to: {output_dir}")
